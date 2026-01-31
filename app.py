import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from prompts import get_medical_prompt
from PIL import Image
import io
import os
import requests
import random
import time
import tempfile

# OCR
try:
    import pytesseract
except ImportError:
    pytesseract = None

try:
    from pdf2image import convert_from_bytes
except ImportError:
    convert_from_bytes = None

# ---------------------------------------------------------
# إعداد الصفحة
# ---------------------------------------------------------
st.set_page_config(page_title="MedMate | رفيقك في الكلية", page_icon="🧬", layout="centered")

st.markdown("""
<style>
/* شيلنا لون الخلفية عشان يظبط أوتوماتيك مع وضع الجهاز */
.stApp { direction: rtl; text-align: right; }

h1, h2, h3, p, div, .stMarkdown, .caption { text-align: right; font-family: 'Segoe UI', Tahoma, Geneva, sans-serif; }
section[data-testid="stSidebar"] { direction: rtl; text-align: right; }
.stTextInput input, .stTextArea textarea, .stSelectbox div[data-baseweb="select"] { direction: rtl; text-align: right; }
.stCheckbox { direction: rtl; text-align: right; }

/* ده زرار التحويل، هنسيبه زي ما هو */
div.stButton > button {
    background-color: #2E86C1; color: white; font-size: 18px; padding: 10px 20px;
    border-radius: 8px; border: none; width: 100%; margin-top: 20px; font-weight: bold;
}

.stAlert { direction: rtl; text-align: right; font-weight: bold; }
#MainMenu {visibility: hidden;}
footer {visibility: hidden !important; height: 0px !important;}
header {visibility: hidden !important;}
div[class^="viewerBadge"] {display: none !important;}
div[class*="viewerBadge"] {display: none !important;}
.stDeployButton {display:none !important;}
[data-testid="stToolbar"] {visibility: hidden !important;}
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------
# أذكار
# ---------------------------------------------------------
AZKAR_LIST = [
    "سبحان الله وبحمده، سبحان الله العظيم 🌿",
    "اللهم صل وسلم وبارك على نبينا محمد ﷺ",
    "لا حول ولا قوة إلا بالله العلي العظيم",
    "أستغفر الله العظيم وأتوب إليه",
    "سبحان الله، والحمد لله، ولا إله إلا الله، والله أكبر",
    "اللهم إنك عفو كريم تحب العفو فاعف عنا",
    "يا حي يا قيوم برحمتك أستغيث",
    "ربّ اشرح لي صدري ويسّر لي أمري"
]

def zikr_update(box, prefix="⏳ جاري المعالجة"):
    box.markdown(f"**{prefix}.. {random.choice(AZKAR_LIST)}** 📿")

# ---------------------------------------------------------
# مفاتيح وأمان
# ---------------------------------------------------------
try:
    GOOGLE_SHEET_URL = st.secrets["GOOGLE_SHEET_URL"]
    api_key = st.secrets["GEMINI_API_KEY"]
except:
    GOOGLE_SHEET_URL = ""
    api_key = None

# ---------------------------------------------------------
# أدوات مساعدة
# ---------------------------------------------------------
def convert_images_to_pdf(image_files):
    images = []
    for file in image_files:
        img = Image.open(file)
        if img.mode != 'RGB':
            img = img.convert('RGB')
        images.append(img)
    if not images:
        return None
    pdf_io = io.BytesIO()
    images[0].save(pdf_io, format='PDF', save_all=True, append_images=images[1:])
    pdf_io.seek(0)
    return pdf_io

# ---------------------------------------------------------
# OCR
# ---------------------------------------------------------
def ocr_image(image):
    if pytesseract is None:
        raise RuntimeError("pytesseract غير مثبت.")
    return pytesseract.image_to_string(image, lang='ara+eng', config='--psm 3')

def process_with_standard_ocr(files, status_box):
    result_text = ""

    for i, f in enumerate(files):
        zikr_update(status_box, "📄 جاري استخراج النص (OCR)")
        time.sleep(0.8)

        if f.type == "application/pdf":
            if convert_from_bytes is None:
                result_text += "\n⚠️ pdf2image غير مثبت لمعالجة PDF.\n"
                continue
            pages = convert_from_bytes(f.getvalue())
            for idx, page in enumerate(pages):
                zikr_update(status_box, f"📄 OCR صفحة {idx+1}")
                text = ocr_image(page)
                result_text += f"\n\n--- صفحة {idx+1} من {f.name} ---\n{text}"
        else:
            img = Image.open(f)
            text = ocr_image(img)
            result_text += f"\n\n--- محتوى الصورة: {f.name} ---\n{text}"

    return result_text

# ---------------------------------------------------------
# Word Formatting
# ---------------------------------------------------------
def add_markdown_paragraph(parent, text, style='Normal', align=None):
    if hasattr(parent, 'add_paragraph'):
        p = parent.add_paragraph(style=style)
    else:
        p = parent

    text = text.replace('***', '**')  # نسيب bold فقط
    if align:
        p.alignment = align
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if any("\u0600" <= c <= "\u06FF" for c in text) else WD_ALIGN_PARAGRAPH.LEFT

    parts = text.split('**')
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True if i % 2 == 1 else False
    return p

def add_page_border(doc):
    sec_pr = doc.sections[0]._sectPr
    pg_borders = OxmlElement('w:pgBorders')
    pg_borders.set(qn('w:offsetFrom'), 'page')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '24')
        border.set(qn('w:color'), 'auto')
        pg_borders.append(border)
    sec_pr.append(pg_borders)

def create_word_table(doc, table_lines):
    if not table_lines:
        return
    cleaned_rows = []
    for line in table_lines:
        if '---' in line:
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        cleaned_rows.append(cells)

    if not cleaned_rows:
        return

    table = doc.add_table(rows=len(cleaned_rows), cols=len(cleaned_rows[0]))
    table.style = 'Table Grid'

    for r_idx, row_data in enumerate(cleaned_rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < len(row.cells):
                cell = row.cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                add_markdown_paragraph(p, cell_text,
                                       align=WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else None)
                if r_idx == 0:
                    for run in p.runs:
                        run.font.bold = True
    doc.add_paragraph("")

def create_styled_word_doc(text_content, user_title):
    doc = Document()
    add_page_border(doc)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    clean_title = user_title.replace('*', '').replace('#', '').strip()
    main_heading = doc.add_heading(clean_title, 0)
    main_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in main_heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

    lines = text_content.split('\n')
    table_buffer = []

    for line in lines:
        line = line.strip()
        if line.startswith('|') and line.endswith('|'):
            table_buffer.append(line)
            continue
        else:
            if table_buffer:
                create_word_table(doc, table_buffer)
                table_buffer = []

        if not line:
            continue

        if line.startswith('#'):
            clean_text = line.lstrip('#').replace('*', '').strip()
            h = doc.add_heading(clean_text, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.RIGHT if any("\u0600" <= c <= "\u06FF" for c in line) else WD_ALIGN_PARAGRAPH.LEFT
            for run in h.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif line.startswith('* ') or line.startswith('- '):
            clean_text = line.lstrip('* ').lstrip('- ').strip()
            add_markdown_paragraph(doc, clean_text, style='List Bullet')

        else:
            add_markdown_paragraph(doc, line)

    if table_buffer:
        create_word_table(doc, table_buffer)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ---------------------------------------------------------
# UI
# ---------------------------------------------------------
import streamlit as st

# شيلنا تحديد اللون الأسود (color: #0e1117) عشان يتلون لوحده حسب الوضع
st.markdown("""
<div style="text-align: right; direction: rtl; font-family: sans-serif;">
    <h1>MedMate | رفيقك الذكي في الكلية 🧬</h1>
    <h3>حوّل صور المحاضرات لملفات Word في ثوانٍ! ⚡</h3>
    <p style="font-size: 1.1em; opacity: 0.8;">
    من أخ لأخيه.. طورنا MedMate علشان يوفرلك وقت للمذاكرة أو العبادة.<br>
    صور المحاضرة، ارفعها هنا، واستلمها ملف Word منسق وجاهز للمذاكرة فورًا.
    </p>
    <small style="opacity: 0.6;">* متاح مجانًا هدية لطلبة طب بني سويف.</small>
</div>
""", unsafe_allow_html=True)

st.divider()

# استخدمنا st.info بدل HTML يدوي عشان تظبط ألوانها أوتوماتيك مع الوضع الليلي والنهاري
st.info("""
**💌 رسالة ودعوة**

العمل ده **صدقة جارية** لدفعة طب بني سويف. لو الأداة فادتك، ادعِ للقائمين عليها بظهر الغيب ❤️
ولو واجهتك مشكلة، ابعتها هنا وهنحلها فورًا بإذن الله.
""")

with st.form(key='feedback_form'):
    # ... باقي الكود زي ما هو ...
    feedback_text = st.text_area("رسالتك:", placeholder="اكتب دعوتك أو اقتراحك هنا...")
    submit_feedback = st.form_submit_button(label='إرسال الرسالة 📨')
    if submit_feedback and feedback_text and GOOGLE_SHEET_URL:
        try:
            requests.post(GOOGLE_SHEET_URL, json={"feedback": feedback_text}, timeout=10)
            st.success("جزاك الله خيرًا! رسالتك وصلت ❤️")
        except:
            st.error("عذرًا، حدث خطأ أثناء الإرسال. لكن نيتك وصلت.❤️")

st.divider()

if 'converted_text' not in st.session_state:
    st.session_state['converted_text'] = ""

uploaded_files = st.file_uploader(
    "📂 ارفع الصور أو ملفات PDF",
    type=['png', 'jpg', 'jpeg', 'pdf'],
    accept_multiple_files=True
)
st.caption("💡 نصيحة أخوية: عشان الموقع يشتغل بسرعة، يفضل ترفع **10-15 صورة** أو **ملف PDF واحد (لا يزيد عن 50 صفحة)** في المرة الواحدة.")

st.write("---")
processing_method = st.radio(
    "⚙️ اختر طريقة المعالجة:",
    ["الذكاء الاصطناعي (AI) - تنسيق ممتاز ✨", "نظام OCR العادي - Tesseract (مجاني بلا حدود) 📄"],
    index=0
)
st.write("---")

# 1. القائمة المنسدلة (خلينا الاختيارات عربي عشان التناسق)
doc_type_selection = st.selectbox(
    "اختار نوع الملف يا دكتور:",
    options=["محاضرات / مذكرات (Notes)", "امتحانات (MCQ / Exam)"],
    index=0
)

# 2. رسائل التوضيح (فصلنا الكود سطور عشان يبقى مقروء ليك مستقبلاً)
if "محاضرات" in doc_type_selection:
    st.info("ℹ️ للمحاضرات والمذكرات: هيتم التنسيق كفقرات وعناوين وشرح متصل.")
elif "امتحانات" in doc_type_selection:
    st.info("ℹ️ للامتحانات: هيتم التنسيق كأسئلة منفصلة واختيارات دقيقة.")

# 3. الأعمدة
col1, col2 = st.columns(2)

# مثال إزاي تملى الأعمدة صح للعربي
with col1:
    is_handwritten = st.checkbox("✍️ هل الملف يحتوى نصوص بخط اليد؟")
with col2:
    user_filename = st.text_input("اسم الملف:", value="MedMate Note")

# ---------------------------------------------------------
# زر التنفيذ
# ---------------------------------------------------------
if st.button("توكلنا على الله.. ابدأ التحويل 🚀"):
    if not uploaded_files:
        st.warning("⚠️ الرجاء رفع الملفات أولاً.")
    elif not api_key and "AI" in processing_method:
        st.error("⚠️ لم يتم العثور على مفتاح API في الإعدادات! يرجى التواصل مع المطور.")
    else:
        status_text = st.empty()
        progress_bar = st.progress(0)

        image_files = [f for f in uploaded_files if f.type.startswith("image/")]
        pdf_files = [f for f in uploaded_files if f.type == "application/pdf"]
        final_content = ""

        # -------------------------------------------------
        # مسار OCR
        # -------------------------------------------------
        if "OCR" in processing_method:
            try:
                final_content = process_with_standard_ocr(uploaded_files, status_text)
                st.session_state['converted_text'] = final_content
                status_text.success("✅ تم استخراج النص بنجاح (OCR)!")
                st.balloons()
            except Exception as e:
                st.error(f"خطأ أثناء OCR: {e}")

        # -------------------------------------------------
        # مسار الذكاء الاصطناعي
        # -------------------------------------------------
        else:
            try:
                genai.configure(api_key=api_key)
                model = genai.GenerativeModel('gemini-flash-latest')

                # ---- دمج الصور في PDF واحد ----
                if image_files:
                    zikr_update(status_text, "📦 جاري دمج الصور")
                    pdf_data = convert_images_to_pdf(image_files)
                    if not pdf_data:
                        raise RuntimeError("فشل دمج الصور.")

                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                        tmp.write(pdf_data.read())
                        temp_name = tmp.name

                    try:
                        zikr_update(status_text, "📤 جاري رفع الملف")
                        g_file = genai.upload_file(temp_name)

                        for _ in range(60):
                            zikr_update(status_text, "⏳ انتظار معالجة الملف")
                            time.sleep(2)
                            g_file = genai.get_file(g_file.name)
                            if g_file.state.name != "PROCESSING":
                                break
                        else:
                            raise TimeoutError("انتهت مهلة معالجة الملف.")

                        zikr_update(status_text, "🧠 جاري التحليل بالذكاء الاصطناعي")
                        response = model.generate_content(
                            [get_medical_prompt(doc_type_selection, is_handwritten), g_file]
                        )
                        final_content += response.text
                        progress_bar.progress(0.5 if pdf_files else 1.0)
                    finally:
                        os.remove(temp_name)

                # ---- معالجة ملفات PDF ----
                for i, pdf in enumerate(pdf_files):
                    zikr_update(status_text, f"📑 جاري تحليل {pdf.name}")
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                        tmp.write(pdf.getvalue())
                        temp_pdf = tmp.name

                    try:
                        g_pdf = genai.upload_file(temp_pdf)
                        for _ in range(60):
                            zikr_update(status_text, "⏳ انتظار معالجة الملف")
                            time.sleep(2)
                            g_pdf = genai.get_file(g_pdf.name)
                            if g_pdf.state.name != "PROCESSING":
                                break
                        else:
                            raise TimeoutError("انتهت مهلة معالجة الملف.")

                        zikr_update(status_text, "🧠 جاري التحليل بالذكاء الاصطناعي")
                        response = model.generate_content(
                            [get_medical_prompt(doc_type_selection, is_handwritten), g_pdf]
                        )
                        final_content += f"\n\nSource: {pdf.name}\n" + response.text
                        progress_bar.progress((i + 1) / len(pdf_files))
                    finally:
                        os.remove(temp_pdf)

                st.session_state['converted_text'] = final_content
                status_text.success("✅ تم التحويل بنجاح يا دكتور!")
                st.balloons()

            # ---- Fallback تلقائي للـ OCR عند نفاذ الرصيد ----
            except Exception as e:
                error_msg = str(e).lower()
                if "429" in error_msg or "quota" in error_msg:
                    st.error("🛑 تم الوصول للحد الأقصى اليومي لاستخدام الذكاء الاصطناعي.")
                    if st.button("اضغط هنا للتحويل باستخدام OCR فورًا 📄"):
                        try:
                            final_content = process_with_standard_ocr(uploaded_files, status_text)
                            st.session_state['converted_text'] = final_content
                            st.rerun()
                        except Exception as ex:
                            st.error(f"فشل OCR الاحتياطي: {ex}")
                else:
                    st.error(f"خطأ تقني: {e}")

# ---------------------------------------------------------
# عرض النتائج
# ---------------------------------------------------------
if st.session_state['converted_text']:
    st.divider()
    docx_file = create_styled_word_doc(st.session_state['converted_text'], user_filename)
    st.success("🎉 اتفضل يا دكتور، ملفك جاهز!")
    st.download_button(
        label=f"💾 تحميل ملف الوورد ({user_filename}.docx)",
        data=docx_file.getvalue(),
        file_name=f"{user_filename}.docx",
        use_container_width=True
    )

    st.subheader("📝 مراجعة النص")
    tab1, tab2 = st.tabs(["✍️ تعديل", "👁️ معاينة"])
    with tab1:
        edited = st.text_area("عدل هنا:", value=st.session_state['converted_text'],
                              height=400, label_visibility="collapsed")
        st.session_state['converted_text'] = edited
    with tab2:
        st.markdown(st.session_state['converted_text'])












